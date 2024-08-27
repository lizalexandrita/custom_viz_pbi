from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('Requirements Document: Creation of ActivityLog Table in TEST_DB1 Database', level=1)

# Project Overview
doc.add_heading('1. Project Overview', level=2)
doc.add_paragraph('Objective: To create a table named ActivityLog in the TEST_DB1 database that tracks user activities '
                  'related to scanning or putting away items within specific locations. '
                  'The table will store details such as the cutoff time for activities, the last scanned time, user and '
                  'location identifiers, item identifiers, and activity details.')

# Table Name
doc.add_heading('2. Table Name', level=2)
doc.add_paragraph('ActivityLog')

# Database
doc.add_heading('3. Database', level=2)
doc.add_paragraph('TEST_DB1')

# Table Schema
doc.add_heading('4. Table Schema', level=2)

# Create a table for the schema
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Column Name'
hdr_cells[1].text = 'Data Type'
hdr_cells[2].text = 'Constraints'
hdr_cells[3].text = 'Description'

# Populate the table schema
schema = [
    ('CutoffTime', 'DATETIME', 'NOT NULL', 'The latest possible time within a day and location for an item to be scanned or put away.'),
    ('LastScanned', 'DATETIME', 'NOT NULL', 'The exact time when the user finished scanning or putting away the item.'),
    ('UserId', 'INT', 'NOT NULL', 'A unique identifier for the user performing the scan or put away activity.'),
    ('LocationId', 'INT', 'NOT NULL', 'A unique identifier for the location where the scanning or putting away is occurring.'),
    ('TenantId', 'INT', 'NOT NULL', 'A unique identifier for the tenant associated with the activity, allowing for multi-tenant data management.'),
    ('LocationItemId', 'INT', 'NOT NULL', 'A unique identifier for the specific item being scanned or put away within the location.'),
    ('CountedQuantity', 'INT', 'NOT NULL', 'The quantity of the item that was actually scanned or put away during the activity.'),
    ('UserActivityId', 'INT', 'NOT NULL', 'A unique identifier that distinguishes whether the activity was a scanning or putting away action.'),
]

for col_name, data_type, constraint, description in schema:
    row_cells = table.add_row().cells
    row_cells[0].text = col_name
    row_cells[1].text = data_type
    row_cells[2].text = constraint
    row_cells[3].text = description

# Constraints
doc.add_heading('5. Constraints', level=2)
doc.add_paragraph('Primary Key: A composite primary key will be created using UserId, LocationId, LocationItemId, and UserActivityId to uniquely identify each record.')
doc.add_paragraph('Foreign Keys:')
doc.add_paragraph('''
- UserId will reference the Users table (if it exists).
- LocationId will reference the Locations table (if it exists).
- TenantId will reference the Tenants table (if it exists).
- LocationItemId will reference the LocationItems table (if it exists).
- UserActivityId will reference the UserActivities table (if it exists).
''')

# Indexes
doc.add_heading('6. Indexes', level=2)
doc.add_paragraph('Indexes for optimization:')
doc.add_paragraph('''
- Index on CutoffTime and LastScanned to optimize queries involving time-based operations.
- Index on LocationId and LocationItemId to optimize location-based queries.
''')

# Data Retention and Archiving
doc.add_heading('7. Data Retention and Archiving', level=2)
doc.add_paragraph('Determine the data retention policy for this table, considering the volume of data and the potential need for historical data analysis. Archiving strategies should be defined if necessary.')

# Security Considerations
doc.add_heading('8. Security Considerations', level=2)
doc.add_paragraph('Ensure appropriate permissions are set on the ActivityLog table, restricting access to sensitive information (such as UserId and TenantId) to authorized users only.')

# Validation and Testing
doc.add_heading('9. Validation and Testing', level=2)
doc.add_paragraph('After creation, validate the table structure by inserting sample data. Test queries to ensure that CutoffTime correctly reflects the maximum time within a day and location, and LastScanned accurately records the activity completion time.')

# Creation Script
doc.add_heading('10. Creation Script', level=2)
doc.add_paragraph('Below is the SQL script for creating the ActivityLog table:')
doc.add_paragraph('''
CREATE TABLE ActivityLog (
    CutoffTime DATETIME NOT NULL,
    LastScanned DATETIME NOT NULL,
    UserId INT NOT NULL,
    LocationId INT NOT NULL,
    TenantId INT NOT NULL,
    LocationItemId INT NOT NULL,
    CountedQuantity INT NOT NULL,
    UserActivityId INT NOT NULL,
    PRIMARY KEY (UserId, LocationId, LocationItemId, UserActivityId),
    FOREIGN KEY (UserId) REFERENCES Users(UserId),  -- If Users table exists
    FOREIGN KEY (LocationId) REFERENCES Locations(LocationId),  -- If Locations table exists
    FOREIGN KEY (TenantId) REFERENCES Tenants(TenantId),  -- If Tenants table exists
    FOREIGN KEY (LocationItemId) REFERENCES LocationItems(LocationItemId),  -- If LocationItems table exists
    FOREIGN KEY (UserActivityId) REFERENCES UserActivities(UserActivityId)  -- If UserActivities table exists
);

-- Indexes for optimization
CREATE INDEX idx_cutoff_time ON ActivityLog (CutoffTime);
CREATE INDEX idx_last_scanned ON ActivityLog (LastScanned);
CREATE INDEX idx_location_item ON ActivityLog (LocationId, LocationItemId);
''')

# Deployment
doc.add_heading('11. Deployment', level=2)
doc.add_paragraph('The creation script should be reviewed and approved by the database administrator (DBA). After approval, deploy the script to the TEST_DB1 database in the development environment for testing.')

# Review and Approval
doc.add_heading('12. Review and Approval', level=2)
doc.add_paragraph('This document and the associated creation script should be reviewed by the project stakeholders, including the DBA and development team, before deployment to ensure all requirements are met.')

# Save the document
doc.save("Requirements_Document_ActivityLog.docx")
